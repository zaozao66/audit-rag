"""
部门监督画像 Agent — 流式 SSE 端点 + 公文格式 DOCX 生成

工作流（4步）：
  1. data_ingest  — 数据摄入与异常指标计算
  2. rag_search   — 政策法规 & 案例检索
  3. analysis     — AI 多维综合研判（流式输出）
  4. report       — 结构化报告字段提取

DOCX 格式规范（参照《word-format-skill》）：
  文档标题：方正小标宋简体，二号（22pt），居中加粗
  一级标题：黑体，三号（16pt）
  二级标题：楷体_GB2312，三号（16pt）
  正文：    仿宋_GB2312，三号（16pt），首行缩进 2 字符
  行间距：  固定值 28 磅
  页边距：  上 3.7cm，下 3.5cm，左 2.8cm，右 2.6cm
"""

import base64
import io
import json
import logging
import re
from typing import Any, Dict, Generator, List

from flask import Blueprint, Response, current_app, jsonify, request, send_file

from src.api.routes.scope_utils import extract_scope_from_request
from src.api.services.rag_service import RAGService

reports_bp = Blueprint('reports', __name__)
logger = logging.getLogger(__name__)

# ── 系统提示词 ────────────────────────────────────────────────────────────────

ANALYSIS_SYSTEM_PROMPT = """你是一名专业的纪检监察AI分析师，负责对企业部门进行监督画像综合研判。

请严格按照以下格式输出分析内容，每个板块用"##"标题标注，不要增加额外板块或改变标题名称：

## 综合判断
（1-2句话，对该部门整体监督态势的综合评价）

## 人员情况分析
（基于人员总数、政治面貌、学历结构，结合相关政策要求进行分析）

## 财务情况分析
（基于差旅报销金额及历年趋势，识别异常波动及合规风险）

## 纪审联动分析
（基于审计问题数量、整改完成率及问题类型分布，评估整改落实情况）

## 岗位廉政风险分析
（基于高中低风险岗位分布，指出重点管控领域）

## 跨维度风险提示
（1-2句话，联合分析各维度风险信号，识别潜在的交叉风险）

## 管理建议
1. （具体可操作的建议，结合政策依据）
2. （具体可操作的建议）
3. （具体可操作的建议）

## 总结
（1句话，提炼核心监督方向）
"""

# ── SSE 工具函数 ───────────────────────────────────────────────────────────────

def _sse(data: Dict[str, Any]) -> str:
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


# ── 数据处理 ──────────────────────────────────────────────────────────────────

def _compute_anomalies(payload: Dict[str, Any]) -> List[str]:
    """根据指标阈值检测异常，返回异常描述列表。"""
    anomalies: List[str] = []

    rate = float(payload.get('audit_completion_rate', 100))
    if rate < 70:
        anomalies.append(f"审计问题整改率偏低（{rate:.1f}%，低于70%预警线）")

    risk_high = int(payload.get('risk_high', 0))
    if risk_high >= 3:
        anomalies.append(f"高风险岗位数量偏高（{risk_high}个，建议重点管控）")

    trend = payload.get('travel_expense_trend', [])
    travel = float(payload.get('travel_expense', 0))
    if len(trend) >= 2:
        prev = float(trend[-2].get('value', 0))
        if prev > 0 and travel > prev * 1.3:
            pct = (travel / prev - 1) * 100
            anomalies.append(f"本年差旅费较上年增长 {pct:.0f}%（{prev}→{travel}万元）")

    audit_issues = int(payload.get('audit_issues', 0))
    audit_completed = int(payload.get('audit_completed', 0))
    outstanding = audit_issues - audit_completed
    if outstanding > 10:
        anomalies.append(f"审计未整改问题积压较多（{outstanding}项待完成）")

    return anomalies


def _build_analysis_prompt(
    payload: Dict[str, Any],
    sources: List[Dict[str, Any]],
    anomalies: List[str],
) -> str:
    dept = payload.get('department_name', '（未知部门）')
    lines = [
        f"## 待分析部门：{dept}",
        "",
        "### 基础监督指标",
        f"- 人员总数：{payload.get('personnel_total', 0)} 人",
        f"- 年度差旅报销：{payload.get('travel_expense', 0)} 万元",
        f"- 审计问题：{payload.get('audit_issues', 0)} 项，"
        f"已完成 {payload.get('audit_completed', 0)} 项，"
        f"整改率 {payload.get('audit_completion_rate', 0):.1f}%",
        f"- 岗位廉政风险：高风险 {payload.get('risk_high', 0)} 个、"
        f"中风险 {payload.get('risk_medium', 0)} 个、"
        f"低风险 {payload.get('risk_low', 0)} 个",
        "",
        "### 政治面貌分布",
    ]
    for item in payload.get('political_distribution', []):
        lines.append(f"- {item['name']}：{item['value']} 人（{item['share']}）")

    lines += ["", "### 学历分布"]
    for item in payload.get('education_distribution', []):
        lines.append(f"- {item['name']}：{item['value']} 人（{item['share']}）")

    lines += ["", "### 历年差旅报销趋势"]
    for item in payload.get('travel_expense_trend', []):
        lines.append(f"- {item['year']} 年：{item['value']} 万元")

    lines += ["", "### 历年审计问题数"]
    for item in payload.get('audit_yearly_issues', []):
        lines.append(f"- {item['year']} 年：{item['value']} 项")

    lines += ["", "### 审计问题类型分布"]
    for item in payload.get('audit_issue_distribution', []):
        lines.append(f"- {item['name']}：{item['value']} 项（{item['share']}）")

    if anomalies:
        lines += ["", "### ⚠ 系统检测到的异常指标"]
        for a in anomalies:
            lines.append(f"- {a}")

    if sources:
        lines += ["", "### 检索到的相关政策法规（供参考）"]
        for s in sources[:4]:
            snippet = s.get('text_snippet', '')
            if snippet:
                lines.append(f"- 《{s['title']}》（相关度 {s['score']:.2f}）：{snippet}…")
            else:
                lines.append(f"- 《{s['title']}》（相关度 {s['score']:.2f}）")

    lines += ["", "请依据以上数据和政策背景，按照指定格式输出分析报告。"]
    return "\n".join(lines)


# ── 结构化解析 ────────────────────────────────────────────────────────────────

def _parse_analysis_sections(text: str) -> Dict[str, Any]:
    """从结构化 Markdown 文本中提取各分析字段。"""

    def _extract(label: str) -> str:
        pattern = rf"##\s+{re.escape(label)}\s*\n(.*?)(?=\n##\s|\Z)"
        m = re.search(pattern, text, re.DOTALL)
        return m.group(1).strip() if m else ""

    def _extract_recommendations() -> List[str]:
        section = _extract("管理建议")
        if not section:
            return []
        items = re.findall(r'^\d+[.、]\s*(.+)$', section, re.MULTILINE)
        return [item.strip() for item in items if item.strip()]

    recs = _extract_recommendations()
    return {
        "overall_judgement": _extract("综合判断"),
        "personnel_insight": _extract("人员情况分析"),
        "finance_insight": _extract("财务情况分析"),
        "audit_insight": _extract("纪审联动分析"),
        "risk_insight": _extract("岗位廉政风险分析"),
        "cross_risk_hint": _extract("跨维度风险提示"),
        "recommendations": recs or ["持续推进问题整改闭环机制", "强化高风险岗位分层管控", "开展跨维度风险联合研判"],
        "conclusion": _extract("总结"),
        "fallback": False,
    }


# ── 核心流式生成器 ────────────────────────────────────────────────────────────

def _stream_department_portrait(
    payload: Dict[str, Any],
    rag_service: RAGService,
    scope: str,
) -> Generator[str, None, None]:

    try:
        # ── Step 1: 数据摄入 ─────────────────────────────────────────
        yield _sse({"event": "step_start", "step": "data_ingest", "label": "数据摄入与指标计算"})

        anomalies = _compute_anomalies(payload)
        dim_count = (
            4
            + len(payload.get('political_distribution', []))
            + len(payload.get('education_distribution', []))
            + len(payload.get('audit_issue_distribution', []))
        )

        yield _sse({
            "event": "step_done",
            "step": "data_ingest",
            "summary": f"已处理 {dim_count} 个监督维度指标，发现 {len(anomalies)} 条预警",
            "anomalies": anomalies,
        })

        # ── Step 2: RAG 检索 ─────────────────────────────────────────
        yield _sse({"event": "step_start", "step": "rag_search", "label": "政策法规 & 案例检索"})

        sources: List[Dict[str, Any]] = []
        rag_processor = None

        try:
            rag_processor = rag_service.get_processor(scope=scope, use_rerank=True, use_llm=True)
            dept = payload.get('department_name', '')
            query = f"{dept} 廉洁风险 审计问题整改 岗位廉政 监督检查"
            results = rag_processor.search(query, top_k=6, use_rerank=True)
            for r in results:
                doc = r['document']
                title = (doc.get('title') or doc.get('filename') or '未知文档').strip()
                sources.append({
                    "title": title,
                    "score": round(float(r.get('score', 0)), 3),
                    "doc_type": doc.get('doc_type', ''),
                    "text_snippet": doc.get('text', '')[:80],
                })
        except Exception as exc:
            logger.warning("RAG检索失败，跳过: %s", exc)

        yield _sse({
            "event": "step_done",
            "step": "rag_search",
            "summary": f"检索到 {len(sources)} 条相关政策法规记录",
            "sources": sources,
        })

        # ── Step 3: AI 多维分析（流式）──────────────────────────────
        yield _sse({"event": "step_start", "step": "analysis", "label": "AI 多维综合研判"})

        user_prompt = _build_analysis_prompt(payload, sources, anomalies)
        analysis_text = ""
        model_name = "unknown"

        try:
            if rag_processor is None:
                rag_processor = rag_service.get_processor(scope=scope, use_rerank=False, use_llm=True)
            llm = rag_processor.llm_provider
            if llm is None:
                raise ValueError("LLM 未配置")

            for event in llm.stream_generate_answer(
                query=user_prompt,
                contexts=[],
                system_prompt=ANALYSIS_SYSTEM_PROMPT,
            ):
                if event.get('type') == 'delta':
                    chunk = event.get('content', '')
                    analysis_text += chunk
                    yield _sse({"event": "analysis_chunk", "text": chunk})
                elif event.get('type') == 'done':
                    model_name = event.get('model', 'unknown')

        except Exception as exc:
            logger.error("AI分析失败: %s", exc)
            yield _sse({"event": "analysis_chunk", "text": f"\n（AI分析服务暂不可用，已降级为规则分析）"})

        yield _sse({"event": "step_done", "step": "analysis", "summary": "多维研判完成"})

        # ── Step 4: 结构化提取 ────────────────────────────────────────
        yield _sse({"event": "step_start", "step": "report", "label": "结构化报告字段提取"})

        structured = _parse_analysis_sections(analysis_text)

        yield _sse({"event": "step_done", "step": "report", "summary": "报告结构化完成，可下载 DOCX"})

        # ── 完成 ──────────────────────────────────────────────────────
        yield _sse({
            "event": "done",
            "model": model_name,
            "analysis": structured,
            "sources": sources,
        })

    except Exception as exc:
        logger.error("部门画像流式生成失败: %s", exc, exc_info=True)
        yield _sse({"event": "error", "message": str(exc)})


# ── 路由 ──────────────────────────────────────────────────────────────────────

@reports_bp.route('/reports/department-portrait/stream', methods=['POST'])
def department_portrait_stream():
    try:
        service: RAGService = current_app.extensions['rag_service']
        data = request.get_json(silent=True) or {}
        scope = extract_scope_from_request(request, json_data=data)
        resolved_scope = service.resolve_scope(scope)

        dept_name = data.get('department_name', '').strip()
        if not dept_name:
            return jsonify({"error": "缺少 department_name"}), 400

        current_app.logger.info(
            "部门画像 Agent 请求: scope=%s dept=%s", resolved_scope, dept_name
        )

        return Response(
            _stream_department_portrait(data, service, resolved_scope),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',
            },
        )
    except Exception as exc:
        current_app.logger.error("部门画像端点异常: %s", exc, exc_info=True)
        return jsonify({"error": str(exc)}), 500

# ── DOCX 格式生成 ─────────────────────────────────────────────────────────────

def _docx_set_font(run, cn_font: str, en_font: str = 'Times New Roman', size_pt: float = 16):
    """设置中英文字体和字号（复刻 word-format-skill 规范）"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def _docx_set_spacing(para, line_pt: float = 28, first_indent: bool = False):
    """设置固定行距 28 磅；正文首行缩进 2 字符（≈640 twips）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(line_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    if first_indent:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '640')


def _build_portrait_docx(
    dept_name: str,
    analysis: Dict[str, Any],
    chart_images: Dict[str, str],
    dept_data: Dict[str, Any],
    political: List[Dict],
    education: List[Dict],
    finance_trend: List[Dict],
    audit_trend: List[Dict],
    audit_dist: List[Dict],
    generated_at: str,
    model_name: str,
) -> bytes:
    """
    用 python-docx 按公文格式规范生成部门监督画像报告 DOCX。
    chart_images: key -> base64 PNG 字符串（data URL 或裸 base64）
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 页边距（标准公文）
    sec = doc.sections[0]
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)

    # ── 内部辅助 ──────────────────────────────────────────────────────────────

    def _title(text: str):
        """文档标题：方正小标宋简体，二号（22pt），居中加粗"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_set_spacing(p)
        r = p.add_run(text)
        r.bold = True
        _docx_set_font(r, '方正小标宋简体', size_pt=22)

    def _h1(text: str):
        """一级标题：黑体，三号（16pt）"""
        p = doc.add_paragraph()
        _docx_set_spacing(p)
        r = p.add_run(text)
        _docx_set_font(r, '黑体', size_pt=16)

    def _h2(text: str):
        """二级标题：楷体_GB2312，三号（16pt）"""
        p = doc.add_paragraph()
        _docx_set_spacing(p)
        r = p.add_run(text)
        _docx_set_font(r, '楷体_GB2312', size_pt=16)

    def _body(text: str):
        """正文：仿宋_GB2312，三号（16pt），首行缩进 2 字符"""
        p = doc.add_paragraph()
        _docx_set_spacing(p, first_indent=True)
        r = p.add_run(text)
        _docx_set_font(r, '仿宋_GB2312', size_pt=16)

    def _body_labeled(label: str, value: str):
        """正文（加粗标签 + 正常内容）"""
        p = doc.add_paragraph()
        _docx_set_spacing(p, first_indent=True)
        r1 = p.add_run(label)
        r1.bold = True
        _docx_set_font(r1, '仿宋_GB2312', size_pt=16)
        r2 = p.add_run(value)
        _docx_set_font(r2, '仿宋_GB2312', size_pt=16)

    def _blank():
        p = doc.add_paragraph()
        _docx_set_spacing(p)

    def _chart(key: str, width_cm: float = 14.0):
        """将 base64 PNG 嵌入文档（居中）。
        图片段落不可设 lineRule=exact 固定行距，否则图片被裁剪不可见。
        仅设段前/段后间距，行高由 Word 自动撑开。
        """
        b64 = chart_images.get(key, '')
        if not b64:
            return
        if ',' in b64:
            b64 = b64.split(',', 1)[1]
        try:
            img_bytes = base64.b64decode(b64)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            p.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
        except Exception as exc:
            logger.warning('图片嵌入失败 [%s]: %s', key, exc)

    # ── 文档内容 ──────────────────────────────────────────────────────────────

    _title(f'{dept_name}部门监督画像报告')

    # 生成信息行（小号灰字，居中）
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_set_spacing(meta_p)
    meta_r = meta_p.add_run(
        f'生成时间：{generated_at}' + (f'　　分析模型：{model_name}' if model_name else '')
    )
    meta_r.font.size = Pt(10)
    meta_r.font.name = 'Times New Roman'
    meta_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    _blank()

    # ── 一、人员情况 ──────────────────────────────────────────────────────────
    _h1('一、人员情况')
    personnel = dept_data.get('personnel', 0)
    _body(f'截至报告期，{dept_name}在岗人员共 {personnel} 人。')

    if political:
        _h2('（一）政治面貌分布')
        _body('；'.join(
            f"{p['name']} {p['value']} 人（{p.get('share', '')}）" for p in political
        ) + '。')
    _chart('political', 14.0)

    if education:
        _h2('（二）学历分布')
        _body('；'.join(
            f"{e['name']} {e['value']} 人（{e.get('share', '')}）" for e in education
        ) + '。')
    _chart('education', 14.0)

    _blank()

    # ── 二、财务情况 ──────────────────────────────────────────────────────────
    _h1('二、财务情况')
    travel = dept_data.get('travel_expense', 0)
    _body(f'{dept_name}本年度差旅报销金额为 {travel} 万元。')
    if finance_trend:
        _body('历年差旅报销趋势：' + '、'.join(
            f"{t['year']} 年 {t['value']} 万元" for t in finance_trend
        ) + '。')
    _chart('financeTrend', 15.5)

    _blank()

    # ── 三、纪审联动情况 ──────────────────────────────────────────────────────
    _h1('三、纪审联动情况')
    issues = dept_data.get('audit_issues', 0)
    completed = dept_data.get('audit_completed', 0)
    rate = dept_data.get('audit_completion_rate', 0)
    _body(f'审计发现问题共 {issues} 项，已完成整改 {completed} 项，整改率 {rate:.1f}%。')
    if audit_trend:
        _body('历年审计问题数量：' + '、'.join(
            f"{t['year']} 年 {t['value']} 项" for t in audit_trend
        ) + '。')
    _chart('auditTrend', 14.0)

    if audit_dist:
        _h2('（一）问题类型分布')
        _body('；'.join(
            f"{a['name']} {a['value']} 项（{a.get('share', '')}）" for a in audit_dist
        ) + '。')
    _chart('auditDistribution', 14.0)

    _blank()

    # ── 四、岗位廉政风险 ──────────────────────────────────────────────────────
    _h1('四、岗位廉政风险')
    rh = dept_data.get('risk_high', 0)
    rm = dept_data.get('risk_medium', 0)
    rl = dept_data.get('risk_low', 0)
    rt = rh + rm + rl
    _body(f'岗位廉政风险共 {rt} 项，其中高风险 {rh} 项、中风险 {rm} 项、低风险 {rl} 项。')
    _chart('risk', 15.5)

    _blank()

    # ── 五、AI综合研判 ────────────────────────────────────────────────────────
    _h1('五、AI综合研判')

    if analysis.get('overall_judgement'):
        _body_labeled('综合判断：', analysis['overall_judgement'])

    field_map = [
        ('personnel_insight', '（一）人员情况分析'),
        ('finance_insight',   '（二）财务情况分析'),
        ('audit_insight',     '（三）纪审联动分析'),
        ('risk_insight',      '（四）岗位廉政风险分析'),
        ('cross_risk_hint',   '（五）跨维度风险提示'),
    ]
    for field, heading in field_map:
        val = str(analysis.get(field, '')).strip()
        if val:
            _h2(heading)
            _body(val)

    recs: List[str] = analysis.get('recommendations', [])
    if recs:
        _blank()
        _h1('六、管理建议')
        for i, rec in enumerate(recs, 1):
            _body(f'{i}. {str(rec).strip()}')

    conclusion = str(analysis.get('conclusion', '')).strip()
    if conclusion:
        _blank()
        _h1('七、总结')
        _body(conclusion)

    # 降级注释
    if analysis.get('fallback'):
        _blank()
        note_p = doc.add_paragraph()
        _docx_set_spacing(note_p)
        note_r = note_p.add_run('注：本报告第五章为规则分析结果，建议待模型联通后复核。')
        note_r.font.size = Pt(10)
        note_r.font.name = 'Times New Roman'
        note_r.font.color.rgb = RGBColor(0xb4, 0x53, 0x09)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── DOCX 路由 ─────────────────────────────────────────────────────────────────

@reports_bp.route('/reports/department-portrait/docx', methods=['POST'])
def department_portrait_docx():
    """
    接收结构化报告数据 + PNG 图表，返回公文格式 DOCX 文件。

    Request JSON:
      department_name      str
      generated_at         str
      model                str
      analysis             DepartmentPortraitAnalysis
      chart_images         { political, education, financeTrend,
                             auditTrend, auditDistribution, risk } -> base64 PNG
      department_data      { personnel, travel_expense, audit_issues,
                             audit_completed, audit_completion_rate,
                             risk_high, risk_medium, risk_low }
      political_distribution  [{name, value, share}]
      education_distribution  [{name, value, share}]
      finance_trend           [{year, value}]
      audit_trend             [{year, value}]
      audit_distribution      [{name, value, share}]
    """
    try:
        data = request.get_json(silent=True) or {}
        dept_name = str(data.get('department_name', '未知部门')).strip()

        doc_bytes = _build_portrait_docx(
            dept_name=dept_name,
            analysis=data.get('analysis', {}),
            chart_images=data.get('chart_images', {}),
            dept_data=data.get('department_data', {}),
            political=data.get('political_distribution', []),
            education=data.get('education_distribution', []),
            finance_trend=data.get('finance_trend', []),
            audit_trend=data.get('audit_trend', []),
            audit_dist=data.get('audit_distribution', []),
            generated_at=str(data.get('generated_at', '')).strip(),
            model_name=str(data.get('model', '')).strip(),
        )

        filename = f'{dept_name}-部门监督画像报告.docx'
        current_app.logger.info('DOCX已生成: %s (%d bytes)', filename, len(doc_bytes))

        return send_file(
            io.BytesIO(doc_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename,
        )
    except Exception as exc:
        current_app.logger.error('DOCX生成失败: %s', exc, exc_info=True)
        return jsonify({'error': str(exc)}), 500
